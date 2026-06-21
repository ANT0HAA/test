"""
База знаний на ChromaDB.
У каждого агента в каждой отрасли — своя коллекция: имя `{industry}_{agent}`.
Документы разбиваются на чанки и хранятся с метаданными.
"""
import os
import re
from pathlib import Path

import chromadb
from chromadb.utils import embedding_functions

from config import settings
from agents.definitions import DEFAULT_INDUSTRY


# ─── Singleton ChromaDB client ─────────────────────────────────────────

_client: chromadb.ClientAPI | None = None


def get_client() -> chromadb.ClientAPI:
    global _client
    if _client is None:
        Path(settings.chroma_path).mkdir(parents=True, exist_ok=True)
        _client = chromadb.PersistentClient(path=settings.chroma_path)
    return _client


def get_collection(agent_name: str, industry: str = DEFAULT_INDUSTRY) -> chromadb.Collection:
    """Получить (или создать) коллекцию агента в рамках отрасли (изоляция по `{industry}_{agent}`)."""
    ef = embedding_functions.DefaultEmbeddingFunction()
    return get_client().get_or_create_collection(
        name=f"{industry}_{agent_name}",
        embedding_function=ef,
        metadata={"agent": agent_name, "industry": industry},
    )


def list_collections() -> list[str]:
    return [c.name for c in get_client().list_collections()]


def get_project_collection(project_id: str) -> chromadb.Collection:
    """Коллекция материалов конкретного проекта (присланные/готовые проекты)."""
    ef = embedding_functions.DefaultEmbeddingFunction()
    return get_client().get_or_create_collection(
        name=f"proj_{project_id}",
        embedding_function=ef,
        metadata={"project_id": project_id},
    )


def delete_project_collection(project_id: str) -> None:
    """Удалить коллекцию материалов проекта (при удалении проекта)."""
    try:
        get_client().delete_collection(f"proj_{project_id}")
    except Exception:
        pass  # коллекции могло не быть


def project_search(query: str, project_id: str, n_results: int = 5) -> str:
    """Поиск по материалам проекта. Пусто, если материалов нет."""
    if not project_id:
        return ""
    col = get_project_collection(project_id)
    if col.count() == 0:
        return ""
    n_results = min(n_results, col.count())
    results = col.query(query_texts=[query], n_results=n_results)
    docs = results.get("documents", [[]])[0]
    metas = results.get("metadatas", [[]])[0]
    if not docs:
        return ""
    return "\n\n---\n\n".join(
        f"[{(m or {}).get('source', '')}]\n{d}" for d, m in zip(docs, metas)
    )


def add_project_text(text: str, project_id: str, filename: str = "исходные данные") -> int:
    """Добавить произвольный текст (напр. заполненную форму) в материалы проекта."""
    chunks = _chunk_text(text)
    if not chunks:
        return 0
    col = get_project_collection(project_id)
    existing = set(col.get()["ids"])
    ids, docs, metas = [], [], []
    for i, chunk in enumerate(chunks):
        doc_id = f"{filename}_{i}"
        if doc_id in existing:
            # перезаписываем «исходные данные» свежими значениями
            col.delete(ids=[doc_id])
        ids.append(doc_id)
        docs.append(chunk)
        metas.append({"source": filename, "chunk": i})
    if ids:
        col.add(ids=ids, documents=docs, metadatas=metas)
    return len(ids)


async def add_project_file(file_path: str, project_id: str) -> int:
    """Добавить файл в материалы проекта (txt/pdf/docx/xlsx). OCR скан-PDF — в потоке."""
    import asyncio
    # Чтение/распознавание может быть тяжёлым (OCR) — не блокируем event loop
    text = await asyncio.to_thread(_read_file_text, Path(file_path))
    chunks = _chunk_text(text)
    if not chunks:
        return 0
    col = get_project_collection(project_id)
    existing = set(col.get()["ids"])
    name = Path(file_path).name
    ids, docs, metas = [], [], []
    for i, chunk in enumerate(chunks):
        doc_id = f"{name}_{i}"
        if doc_id in existing:
            continue
        ids.append(doc_id)
        docs.append(chunk)
        metas.append({"source": name, "chunk": i})
    if ids:
        col.add(ids=ids, documents=docs, metadatas=metas)
    return len(ids)


# ─── Управление материалами проекта (просмотр/правка/удаление) ─────────

def list_project_materials(project_id: str) -> list[dict]:
    """Все фрагменты материалов проекта: [{id, source, chunk, text}], по источникам."""
    col = get_project_collection(project_id)
    if col.count() == 0:
        return []
    data = col.get(include=["documents", "metadatas"])
    ids = data.get("ids", []) or []
    docs = data.get("documents", []) or []
    metas = data.get("metadatas", []) or []
    items = [
        {"id": i, "source": (m or {}).get("source", ""),
         "chunk": (m or {}).get("chunk", 0), "text": d}
        for i, d, m in zip(ids, docs, metas)
    ]
    items.sort(key=lambda x: (x["source"], x["chunk"]))
    return items


def update_project_material(project_id: str, frag_id: str, text: str) -> bool:
    """Изменить текст фрагмента материала проекта."""
    col = get_project_collection(project_id)
    found = col.get(ids=[frag_id])
    if not found.get("ids"):
        return False
    meta = (found.get("metadatas") or [{}])[0] or {}
    col.update(ids=[frag_id], documents=[text], metadatas=[meta])
    return True


def delete_project_material(project_id: str, frag_id: str) -> bool:
    """Удалить один фрагмент материала проекта."""
    col = get_project_collection(project_id)
    if not col.get(ids=[frag_id]).get("ids"):
        return False
    col.delete(ids=[frag_id])
    return True


def delete_project_source(project_id: str, source: str) -> int:
    """Удалить все фрагменты одного источника (файла). Возвращает число удалённых."""
    col = get_project_collection(project_id)
    found = col.get(where={"source": source})
    ids = found.get("ids", []) or []
    if ids:
        col.delete(ids=ids)
    return len(ids)


def collection_stats(agent_name: str, industry: str = DEFAULT_INDUSTRY) -> dict:
    """Статистика базы знаний агента: число фрагментов и список файлов-источников."""
    col = get_collection(agent_name, industry)
    count = col.count()
    if count == 0:
        return {"chunks": 0, "sources": []}
    metas = col.get(include=["metadatas"]).get("metadatas", []) or []
    sources: list[str] = []
    for m in metas:
        src = (m or {}).get("source")
        if src and src not in sources:
            sources.append(src)
    return {"chunks": count, "sources": sorted(sources)}


# ─── Chunking ──────────────────────────────────────────────────────────

def _chunk_text(text: str, chunk_size: int = 1000, overlap: int = 150) -> list[str]:
    """Разбить текст на перекрывающиеся чанки."""
    text = re.sub(r"\n{3,}", "\n\n", text.strip())
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start += chunk_size - overlap
    return chunks


# ─── Add documents ─────────────────────────────────────────────────────

def add_text(
    text: str,
    agent_name: str,
    filename: str = "manual",
    extra_metadata: dict | None = None,
    industry: str = DEFAULT_INDUSTRY,
) -> int:
    """Добавить текст в базу знаний агента отрасли. Возвращает кол-во добавленных чанков."""
    chunks = _chunk_text(text)
    if not chunks:
        return 0

    col = get_collection(agent_name, industry)
    existing_ids = set(col.get()["ids"])

    ids, documents, metadatas = [], [], []
    for i, chunk in enumerate(chunks):
        doc_id = f"{filename}_{i}"
        if doc_id in existing_ids:
            continue  # пропускаем дубликаты
        ids.append(doc_id)
        documents.append(chunk)
        metadatas.append({"source": filename, "chunk": i, **(extra_metadata or {})})

    if ids:
        col.add(ids=ids, documents=documents, metadatas=metadatas)

    return len(ids)


def _read_file_text(path: Path) -> str:
    """Извлечь текст из файла (txt/pdf/docx/xlsx). Бросает ValueError для .doc и прочих."""
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")

    if suffix == ".pdf":
        try:
            import PyPDF2
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            raise ValueError(f"Ошибка чтения PDF: {e}") from e
        # Скан-PDF без текстового слоя → OCR (если доступен Tesseract)
        if len(text.strip()) < 50:
            from .ocr import ocr_available, ocr_pdf
            if ocr_available():
                ocr_text = ocr_pdf(path)
                if len(ocr_text.strip()) >= 50:
                    return ocr_text
                raise ValueError(
                    "PDF без текстового слоя: OCR не распознал текст "
                    "(проверьте качество скана).")
            raise ValueError(
                "PDF без текстового слоя (скан). Для распознавания нужен Tesseract "
                "(см. backend/.tessdata, config.tesseract_cmd) либо пришлите PDF с текстом.")
        return text

    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception as e:
            raise ValueError(f"Ошибка чтения DOCX: {e}") from e

    if suffix == ".doc":
        raise ValueError("Формат .doc (старый Word) не поддерживается — пересохраните в .docx")

    if suffix in (".xlsx", ".xlsm"):
        try:
            from openpyxl import load_workbook
            wb = load_workbook(str(path), read_only=True, data_only=True)
            try:
                lines: list[str] = []
                for ws in wb.worksheets:
                    lines.append(f"[Лист: {ws.title}]")
                    for row in ws.iter_rows(values_only=True):
                        cells = [str(c).strip() for c in row if c not in (None, "")]
                        if cells:
                            lines.append(" | ".join(cells))
                return "\n".join(lines)
            finally:
                wb.close()
        except Exception as e:
            raise ValueError(f"Ошибка чтения XLSX: {e}") from e

    raise ValueError(f"Неподдерживаемый тип файла: {suffix}")


async def add_file(file_path: str, agent_name: str, industry: str = DEFAULT_INDUSTRY) -> int:
    """Прочитать файл (txt/pdf/docx/xlsx) и добавить в базу агента отрасли."""
    path = Path(file_path)
    return add_text(_read_file_text(path), agent_name, filename=path.name, industry=industry)


# ─── Search ────────────────────────────────────────────────────────────

def search(query: str, agent_name: str, industry: str = DEFAULT_INDUSTRY, n_results: int = 5) -> str:
    """
    Семантический поиск по базе знаний агента отрасли.
    Возвращает конкатенированные релевантные чанки или пустую строку.
    """
    col = get_collection(agent_name, industry)

    # Если коллекция пустая — вернуть пустоту без ошибки
    if col.count() == 0:
        return ""

    n_results = min(n_results, col.count())
    results = col.query(query_texts=[query], n_results=n_results)

    docs = results.get("documents", [[]])[0]
    metas = results.get("metadatas", [[]])[0]

    if not docs:
        return ""

    parts = []
    for doc, meta in zip(docs, metas):
        source = meta.get("source", "")
        parts.append(f"[{source}]\n{doc}")

    return "\n\n---\n\n".join(parts)
