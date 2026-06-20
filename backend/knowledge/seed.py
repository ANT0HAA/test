"""
Наполнение базы знаний агентов документами проекта (batch-ингест).

Берёт папку с документами и раскладывает их по коллекциям агентов отрасли
согласно маппингу «имя файла → агент(ы)». Поддерживает txt / pdf / docx / xlsx.
Сканы без текстового слоя и архивы (.rar/.7z/.doc) пропускаются с пояснением.

Запуск:
    cd backend && .venv\\Scripts\\activate
    python -m knowledge.seed "C:\\путь\\к\\Документы" [--industry ceramics] [--max-pdf-pages 0]

--max-pdf-pages 0 — без ограничения (весь документ); иначе — первые N страниц.
"""
import argparse
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from agents.definitions import DEFAULT_INDUSTRY, get_agents
from knowledge.chroma import add_text


# Маппинг: подстрока имени файла (в нижнем регистре) → агенты-получатели.
# Файл попадает в базу знаний КАЖДОГО указанного агента.
MAPPING: dict[str, list[str]] = {
    "справочник по керамике": ["technologist"],
    "расчет сушил": ["technologist"],
    "варианты расчета сушил": ["technologist"],
    "расчет состава шихты": ["technologist"],
    "регламент": ["technologist"],
    "схема_потоков": ["technologist"],
    "оборудование": ["mechanic", "estimator"],
    "штатное расписание": ["estimator", "documentalist"],
    "теплотехник": ["documentalist"],
    "тэо": ["documentalist"],
    "пояснительная записка": ["documentalist"],
    "геолог": ["technologist"],
    "карты залеган": ["technologist"],
    # ключи по именам архивов (наследуются вложенными файлами без своего маппинга)
    "книги по производству": ["technologist"],
    "лышня": ["technologist"],
    "статья": ["technologist"],
    "прайс": ["estimator"],
}

_TEXT_EXT = {".txt", ".pdf", ".docx", ".xlsx", ".xlsm"}
_ARCHIVE_EXT = {".rar", ".7z", ".zip"}


def _find_7z() -> str | None:
    """Найти 7z.exe (для распаковки .rar/.7z/.zip)."""
    env = os.environ.get("SEVENZIP_PATH")
    if env and Path(env).exists():
        return env
    found = shutil.which("7z") or shutil.which("7za")
    if found:
        return found
    for cand in (r"C:\Program Files\7-Zip\7z.exe", r"C:\Program Files (x86)\7-Zip\7z.exe"):
        if Path(cand).exists():
            return cand
    return None


def _match_agents(filename: str) -> list[str]:
    low = filename.lower()
    agents: list[str] = []
    for key, targets in MAPPING.items():
        if key in low:
            for t in targets:
                if t not in agents:
                    agents.append(t)
    return agents


def _pdf_text(path: Path, max_pages: int) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(str(path))
    pages = reader.pages if max_pages <= 0 else reader.pages[:max_pages]
    return "\n".join((p.extract_text() or "") for p in pages)


def _docx_text(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _xlsx_text(path: Path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(str(path), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        lines.append(f"[Лист: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c not in (None, "")]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract(path: Path, max_pdf_pages: int) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".pdf":
        return _pdf_text(path, max_pdf_pages)
    if suffix == ".docx":
        return _docx_text(path)
    if suffix in (".xlsx", ".xlsm"):
        return _xlsx_text(path)
    return ""


def _ingest_file(path: Path, industry: str, known_agents: set[str],
                 max_pdf_pages: int, label: str, fallback_agents: list[str] | None = None) -> int:
    """Ингест одного файла. Возвращает число добавленных чанков."""
    agents = _match_agents(path.name) or (fallback_agents or [])
    agents = [a for a in agents if a in known_agents]
    if not agents:
        print(f"  — {label}: пропуск (нет маппинга/агенты не из отрасли)")
        return 0
    if path.suffix.lower() not in _TEXT_EXT:
        print(f"  — {label}: пропуск (формат {path.suffix} не поддерживается)")
        return 0
    try:
        text = _extract(path, max_pdf_pages)
    except Exception as e:
        print(f"  ✗ {label}: ошибка чтения — {type(e).__name__}: {e}")
        return 0
    if len(text.strip()) < 50:
        print(f"  — {label}: пропуск (нет текстового слоя — вероятно скан, нужен OCR)")
        return 0

    total = 0
    for agent in agents:
        added = add_text(text, agent, filename=path.name, industry=industry)
        total += added
        print(f"  ✓ {label} → {agent}: +{added} чанков")
    return total


def _ingest_archive(path: Path, industry: str, known_agents: set[str],
                    max_pdf_pages: int, sevenzip: str | None) -> int:
    """Распаковать архив (7z) и заингестить вложенные поддерживаемые файлы."""
    if not sevenzip:
        print(f"  — {path.name}: пропуск (нет 7-Zip; задайте SEVENZIP_PATH)")
        return 0
    archive_agents = _match_agents(path.name)  # наследуется вложенными без своего маппинга
    tmp = Path(tempfile.mkdtemp(prefix="seed_arc_"))
    try:
        r = subprocess.run([sevenzip, "x", "-y", f"-o{tmp}", str(path)],
                           capture_output=True, text=True)
        if r.returncode != 0:
            print(f"  ✗ {path.name}: ошибка распаковки 7z ({r.returncode})")
            return 0
        total = 0
        inner = [p for p in tmp.rglob("*") if p.is_file()]
        print(f"  ▷ {path.name}: распаковано файлов — {len(inner)}")
        for f in inner:
            if f.suffix.lower() not in _TEXT_EXT:
                continue
            total += _ingest_file(f, industry, known_agents, max_pdf_pages,
                                  label=f"{path.name} ▸ {f.name}", fallback_agents=archive_agents)
        return total
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def seed_directory(directory: str, industry: str = DEFAULT_INDUSTRY, max_pdf_pages: int = 0) -> None:
    root = Path(directory)
    if not root.is_dir():
        print(f"Папка не найдена: {root}")
        sys.exit(1)

    known_agents = set(get_agents(industry))
    sevenzip = _find_7z()
    total_chunks = 0
    print(f"Ингест в отрасль «{industry}». Папка: {root}")
    print(f"7-Zip: {sevenzip or 'не найден (архивы будут пропущены)'}\n")

    for path in sorted(root.iterdir()):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()
        if suffix in _ARCHIVE_EXT:
            total_chunks += _ingest_archive(path, industry, known_agents, max_pdf_pages, sevenzip)
        else:
            total_chunks += _ingest_file(path, industry, known_agents, max_pdf_pages, label=path.name)

    print(f"\nГотово. Всего добавлено чанков: {total_chunks}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Наполнение базы знаний агентов документами")
    parser.add_argument("directory", help="Папка с документами")
    parser.add_argument("--industry", default=DEFAULT_INDUSTRY, help="Отрасль (по умолчанию ceramics)")
    parser.add_argument("--max-pdf-pages", type=int, default=0,
                        help="Лимит страниц PDF (0 — без ограничения)")
    args = parser.parse_args()
    seed_directory(args.directory, args.industry, args.max_pdf_pages)


if __name__ == "__main__":
    main()
