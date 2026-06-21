"""
Тесты генераторов отчётов (DOCX / XLSX / PDF) из ProjectExportData.
БД не нужна — данные конструируются вручную.
"""
import io
from datetime import datetime, timezone

import pytest

from export.common import ProjectExportData, AgentSection
from export.docx_report import build_explanatory_note
from export.pdf_report import build_summary_report
from export.xlsx_report import build_equipment_sheet, _extract_heuristic, EquipmentRow


def _data():
    return ProjectExportData(
        project_id="p1", name="Тестовый завод", industry="ceramics",
        created_at=datetime(2026, 6, 20, tzinfo=timezone.utc),
        tasks=["Спроектировать обжигательный корпус"],
        sections=[
            AgentSection("builder", "Конструктор", "Каркас стальной, пролёт 24 м."),
            AgentSection("mechanic", "Инженер-механик",
                         "Туннельная печь — 1 шт.\nВагонетки — 45 шт."),
        ],
        raw_text="Туннельная печь — 1 шт.\nВагонетки — 45 шт.\nДымосос ДН-12 — 2 ед.",
        calc_summary="Производственная программа:\n  • выпуск: 15 000 000 шт/год",
        equipment=[{"role": "Формование", "name": "Пресс СМК-133", "capacity": "20 т/ч", "qty": 1},
                   {"role": "Обжиг", "name": "Туннельная печь", "capacity": "L≈80 м", "qty": 1}],
    )


async def test_docx_valid_and_structured():
    content = await build_explanatory_note(_data())
    from docx import Document
    doc = Document(io.BytesIO(content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА" in text
    assert "Конструктор" in text
    assert "Расчётные данные" in text          # раздел из расчётов
    assert "15 000 000" in text                # цифра из расчётов
    assert len(doc.tables) >= 1  # штамп


async def test_xlsx_equipment_from_calc():
    content = await build_equipment_sheet(_data())
    from openpyxl import load_workbook
    ws = load_workbook(io.BytesIO(content)).active
    cells = [ws.cell(row=r, column=2).value for r in range(5, 12)]
    assert any(c and "СМК-133" in str(c) for c in cells)   # подобранное оборудование


async def test_xlsx_columns_and_rows():
    content = await build_equipment_sheet(_data())
    from openpyxl import load_workbook
    ws = load_workbook(io.BytesIO(content)).active
    headers = [ws.cell(row=4, column=c).value for c in range(1, 7)]
    assert headers == ["№ п/п", "Наименование", "Марка/Тип", "Кол-во", "Ед.", "Примечание"]


async def test_pdf_valid():
    content = await build_summary_report(_data())
    assert content[:5] == b"%PDF-"
    from PyPDF2 import PdfReader
    assert len(PdfReader(io.BytesIO(content)).pages) >= 1


def test_equipment_heuristic_extracts_counts():
    rows = _extract_heuristic(_data())
    assert any(isinstance(r, EquipmentRow) for r in rows)
    # «Туннельная печь — 1 шт.» / «Вагонетки — 45 шт.» — счётные единицы
    names = " ".join(r.name.lower() for r in rows)
    assert "печь" in names or "вагонетки" in names
