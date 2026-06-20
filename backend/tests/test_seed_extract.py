"""
Тесты извлечения текста и маппинга в knowledge/seed (без ChromaDB).
"""
import tempfile
from pathlib import Path

from openpyxl import Workbook
from docx import Document

from knowledge import seed


def test_match_agents_by_filename():
    assert seed._match_agents("Справочник по керамике.pdf") == ["technologist"]
    assert set(seed._match_agents("ТТТ Оборудование ЦЕНА.xlsx")) == {"mechanic", "estimator"}
    assert seed._match_agents("Штатное расписание.xlsx") == ["estimator", "documentalist"]
    assert seed._match_agents("случайный файл.txt") == []


def test_xlsx_text():
    with tempfile.TemporaryDirectory() as td:
        wb = Workbook()
        ws = wb.active
        ws.title = "Оборудование"
        ws.append(["Наименование", "Кол-во"])
        ws.append(["Глинорыхлитель СМК 496", 1])
        p = Path(td) / "eq.xlsx"
        wb.save(p)
        text = seed._xlsx_text(p)
    assert "[Лист: Оборудование]" in text
    assert "Глинорыхлитель СМК 496" in text


def test_docx_text():
    with tempfile.TemporaryDirectory() as td:
        doc = Document()
        doc.add_paragraph("Расчёт состава шихты")
        t = doc.add_table(rows=1, cols=2)
        t.rows[0].cells[0].text = "глина"
        t.rows[0].cells[1].text = "30%"
        p = Path(td) / "calc.docx"
        doc.save(p)
        text = seed._docx_text(p)
    assert "Расчёт состава шихты" in text
    assert "глина | 30%" in text


def test_extract_dispatch_txt():
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "note.txt"
        p.write_text("привет мир", encoding="utf-8")
        assert seed._extract(p, 0) == "привет мир"


def test_doc_routed_to_word(monkeypatch):
    # .doc должен идти через Word-конвертер (на машине без Word — мягкий пропуск)
    assert ".doc" in seed._TEXT_EXT
    monkeypatch.setattr(seed, "_doc_text_via_word", lambda p: "текст из word")
    assert seed._extract(Path("Регламент.doc"), 0) == "текст из word"

