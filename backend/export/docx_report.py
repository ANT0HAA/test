"""
Пояснительная записка → DOCX (python-docx).

Структура: титульный лист с основной надписью (штампом), исходное задание,
разделы по специалистам. Оформление минимальное по ЕСКД/СПДС.
"""
import io

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .common import ProjectExportData, ORG_NAME, ORG_SUBTITLE, lab_lines, balance_lines, firing_lines


def _add_stamp(doc: Document, data: ProjectExportData) -> None:
    """Основная надпись (упрощённый штамп) таблицей."""
    rows = [
        ("Организация", ORG_NAME),
        ("Проект", data.name),
        ("Отрасль", data.industry_name),
        ("Документ", "Пояснительная записка"),
        ("Дата", data.date_str),
        ("Разработал", "Главный конструктор (оркестратор) + специалисты"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = label
        c1.text = value
        c0.paragraphs[0].runs[0].bold = True


async def build_explanatory_note(data: ProjectExportData) -> bytes:
    """Сформировать пояснительную записку и вернуть DOCX как bytes."""
    doc = Document()

    # Базовый шрифт документа
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # ── Титульный лист ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(ORG_NAME)
    run.bold = True
    run.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(ORG_SUBTITLE).italic = True

    doc.add_paragraph()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = h.add_run("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА")
    hr.bold = True
    hr.font.size = Pt(16)

    pname = doc.add_paragraph()
    pname.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pname.add_run(f"по проекту «{data.name}»").font.size = Pt(13)

    doc.add_paragraph()
    _add_stamp(doc, data)

    doc.add_page_break()

    # ── Исходное задание ──
    doc.add_heading("1. Исходное задание", level=1)
    if data.tasks:
        for i, task in enumerate(data.tasks, 1):
            doc.add_paragraph(f"{i}. {task}")
    else:
        doc.add_paragraph("Задание не зафиксировано.")

    # ── Расчётные данные ──
    no = 1
    if data.calc_summary:
        no += 1
        doc.add_heading(f"{no}. Расчётные данные", level=1)
        for para in data.calc_summary.split("\n"):
            para = para.rstrip()
            if para:
                doc.add_paragraph(para)

    # ── Режим обжига ──
    fir_block = firing_lines(data.spec)
    if fir_block:
        no += 1
        doc.add_heading(f"{no}. Режим обжига", level=1)
        for line in fir_block:
            doc.add_paragraph(line)

    # ── Материальный баланс по переделам ──
    bal_block = balance_lines(data.spec)
    if bal_block:
        no += 1
        doc.add_heading(f"{no}. Материальный баланс по переделам", level=1)
        for line in bal_block:
            doc.add_paragraph(line)

    # ── Лаборатория (сырьё и шихта) ──
    lab_block = lab_lines(data.lab)
    if lab_block:
        no += 1
        doc.add_heading(f"{no}. Лаборатория · сырьё и шихта", level=1)
        for line in lab_block:
            doc.add_paragraph(line)

    # ── Разделы по специалистам ──
    section_no = no + 1
    doc.add_heading(f"{section_no}. Проектные решения", level=1)
    if not data.sections:
        doc.add_paragraph("По проекту пока нет проработанных решений специалистов.")
    for idx, section in enumerate(data.sections, 1):
        doc.add_heading(f"{section_no}.{idx}. {section.display_name}", level=2)
        for para in (section.content or "—").split("\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
